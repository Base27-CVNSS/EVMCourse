#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from lxml import html as lxml_html
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "build"
DOWNLOADS = ROOT / "downloads"
ASSETS = BUILD / "book_assets"
DOWNLOADS.mkdir(exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, "/root/.codex/skills/builtins/documents/scripts")
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402

COURSE = json.loads((BUILD / "course.json").read_text(encoding="utf-8"))
INK = RGBColor(0x14, 0x25, 0x1F)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5F, 0x70, 0x69)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"


def clean(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text.replace("–", "-").replace("—", "-").replace("‑", "-").replace("−", "-")


def set_run(run, *, size=None, bold=None, italic=None, color=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str, border: str = "2E74B5"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), border)
    borders.append(left)
    ppr.append(borders)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_repeat_table_header_and_style(table):
    repeat_table_header(table.rows[0])
    for idx, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                set_run(run, size=9.5, bold=True, color=INK)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, size=9.5, color=INK)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = clean(text)
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def set_last_image_alt(doc: Document, description: str):
    shape = doc.inline_shapes[-1]
    shape._inline.docPr.set("descr", clean(description))
    shape._inline.docPr.set("title", clean(description))


def create_numbering(doc: Document, *, bullet=False) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    ppr.extend([tabs, ind])
    lvl.extend([start, fmt, lvl_text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_items(doc: Document, items):
    num_id = create_numbering(doc, bullet=False)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        set_run(p.add_run(clean(text)), color=INK)
    return num_id


def add_bullet_items(doc: Document, items, *, size=11):
    num_id = create_numbering(doc, bullet=True)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if size < 11 else 4)
        p.paragraph_format.line_spacing = 1.05 if size < 11 else 1.208
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        set_run(p.add_run(clean(text)), size=size, color=INK)
    return num_id


def start_on_new_page(paragraph):
    """Start content on a fresh page without creating an empty trailing page."""
    paragraph.paragraph_format.page_break_before = True
    return paragraph


def configure_document(doc: Document, preset: str, running_label: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8 if preset == "narrative_proposal" else 6)
    pf.line_spacing = 1.333 if preset == "narrative_proposal" else 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative_proposal" else WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12 if preset == "narrative_proposal" else 14, 6 if preset == "narrative_proposal" else 7),
        ("Heading 3", 12, DARK_BLUE, 8 if preset == "narrative_proposal" else 10, 4 if preset == "narrative_proposal" else 5),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, left, hanging, after, spacing in [
        ("List Bullet", 0.375, 0.194, 4, 1.208 if preset == "narrative_proposal" else 1.25),
        ("List Number", 0.375, 0.194, 4, 1.208 if preset == "narrative_proposal" else 1.25),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(-hanging)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run(hp.add_run(running_label), size=8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    set_run(fp.add_run("Long Ngo  |  "), size=8.5, color=MUTED)
    add_page_number(fp)

    doc.core_properties.title = COURSE["meta"]["title"]
    doc.core_properties.author = "Long Ngo"
    doc.core_properties.subject = "Environmental systems, AI, LONG, Digital Twin"
    doc.core_properties.keywords = "Environmental Systems; GeoAI; LONG; Carbon; Digital Twin"


def make_diagrams():
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if not Path(bold_path).exists():
        bold_path = font_path
    font = ImageFont.truetype(font_path, 25)
    small = ImageFont.truetype(font_path, 20)
    bold = ImageFont.truetype(bold_path, 27)
    bg = "#f7f9f8"
    navy = "#1f4d78"
    green = "#0b5d4b"
    gold = "#d78b38"

    def box(draw, xy, title, body, fill):
        draw.rounded_rectangle(xy, radius=22, fill=fill, outline="#cbd7d1", width=2)
        x1, y1, x2, y2 = xy
        draw.text((x1 + 22, y1 + 18), title, font=bold, fill=navy)
        draw.multiline_text((x1 + 22, y1 + 58), body, font=small, fill="#32433d", spacing=7)

    im = Image.new("RGB", (1500, 740), bg)
    d = ImageDraw.Draw(im)
    d.text((750, 55), "MÔ HÌNH LONG", font=ImageFont.truetype(bold_path, 42), fill=navy, anchor="mm")
    box(d, (70, 150, 720, 350), "L - Learning & Life-cycle", "Dữ liệu, QA/QC, LCA, bằng chứng\nvà khả năng tổng quát hóa", "#e8f3ef")
    box(d, (780, 150, 1430, 350), "O - Optimization", "Mục tiêu, ràng buộc, Pareto,\nchi phí và phương án khả thi", "#eef3f8")
    box(d, (70, 420, 720, 620), "N - Network", "Mạng sông, sinh thái, hạ tầng,\nchuỗi cung ứng và thể chế", "#fff4e6")
    box(d, (780, 420, 1430, 620), "G - Governance", "Trách nhiệm, pháp lý, tham gia,\ngiám sát con người và fallback", "#f2eef8")
    im.save(ASSETS / "long_framework.png", dpi=(180, 180))

    im = Image.new("RGB", (1600, 560), bg)
    d = ImageDraw.Draw(im)
    labels = [
        ("ĐẦU VÀO", "Vật chất\nNăng lượng\nThông tin"),
        ("TIẾN TRÌNH", "Bảo toàn\nPhản ứng\nĐiều khiển"),
        ("ĐẦU RA", "Tải lượng\nTrạng thái\nDịch vụ"),
        ("PHẢN HỒI", "Cảnh báo\nĐiều chỉnh\nHọc thích ứng"),
    ]
    xs = [40, 430, 820, 1210]
    for i, ((title, body), x) in enumerate(zip(labels, xs)):
        box(d, (x, 140, x + 340, 390), title, body, ["#e8f3ef", "#eef3f8", "#fff4e6", "#f2eef8"][i])
        if i < 3:
            d.line((x + 340, 265, x + 390, 265), fill=green, width=7)
            d.polygon([(x + 390, 265), (x + 370, 251), (x + 370, 279)], fill=green)
    d.arc((170, 380, 1430, 520), 0, 180, fill=gold, width=6)
    d.polygon([(176, 451), (196, 437), (196, 465)], fill=gold)
    im.save(ASSETS / "ipof_loop.png", dpi=(180, 180))

    im = Image.new("RGB", (1500, 760), bg)
    d = ImageDraw.Draw(im)
    layers = [
        ("1. Quan trắc", "IoT - trạm chuẩn - radar - vệ tinh", "#e8f3ef"),
        ("2. Nền dữ liệu", "QA/QC - schema - provenance - phân quyền", "#eef3f8"),
        ("3. Twin core", "Trạng thái - mô hình - đồng hóa - ensemble", "#fff4e6"),
        ("4. Quyết định", "Kịch bản - cảnh báo - tối ưu - phản hồi", "#f2eef8"),
    ]
    y = 65
    for idx, (title, body, fill) in enumerate(layers):
        box(d, (130, y, 1370, y + 125), title, body, fill)
        if idx < len(layers) - 1:
            d.line((750, y + 125, 750, y + 160), fill=green, width=7)
            d.polygon([(750, y + 160), (736, y + 140), (764, y + 140)], fill=green)
        y += 165
    im.save(ASSETS / "digital_twin_architecture.png", dpi=(180, 180))


def add_cover(doc: Document, *, instructor=False):
    for _ in range(5 if not instructor else 3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("GIÁO TRÌNH ĐẠI HỌC MỞ" if not instructor else "SỔ TAY GIẢNG VIÊN"), size=11, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("PHÂN TÍCH HỆ THỐNG MÔI TRƯỜNG\nDỰA TRÊN AI VÀ MÔ HÌNH LONG"), size=29 if not instructor else 25, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("Khoa học hệ thống - GeoAI - Carbon - Công bằng môi trường - Digital Twin"), size=13, italic=True, color=MUTED)
    doc.add_paragraph()
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(author.add_run("Chủ biên: Long Ngo"), size=12, bold=True, color=INK)
    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(edition.add_run("Ấn bản đại học 2.0 - 2026 - 4 tín chỉ"), size=10, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    set_run(note.add_run("Giấy phép MIT cho nội dung gốc và mã nguồn.\nNguồn bên thứ ba giữ nguyên điều kiện sử dụng."), size=9, color=MUTED)
    doc.add_page_break()


def add_callout(doc, label: str, body: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(clean(label) + ". "), bold=True, color=DARK_BLUE)
    set_run(p.add_run(clean(body)), color=INK)
    shade_paragraph(p, LIGHT_GRAY)
    return p


def add_table(doc, rows: list[list[str]], weights=None, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(clean(caption)), size=9.5, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = clean(value)
    widths = column_widths_from_weights(weights or [1] * len(rows[0]), 9360)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120, cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120})
    set_repeat_table_header_and_style(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_formula(doc, formula: str, note: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(clean(formula)), size=13, bold=True, color=DARK_BLUE, font="Cambria Math")
    shade_paragraph(p, "F6F8FA", "D9E2E9")
    if note:
        n = doc.add_paragraph()
        n.paragraph_format.space_before = Pt(0)
        n.paragraph_format.space_after = Pt(8)
        set_run(n.add_run(clean(note)), size=9, italic=True, color=MUTED)


def html_fragment_to_doc(doc, fragment: str):
    wrapper = lxml_html.fragment_fromstring(fragment, create_parent="div")
    for node in wrapper:
        tag = node.tag.lower() if isinstance(node.tag, str) else ""
        if tag == "p":
            p = doc.add_paragraph(clean(node.text_content()))
            for run in p.runs:
                set_run(run)
        elif tag == "h3":
            doc.add_heading(clean(node.text_content()), level=3)
        elif tag == "table":
            caption = ""
            cap = node.find("caption")
            if cap is not None:
                caption = cap.text_content()
            rows = []
            for tr in node.xpath(".//tr"):
                cells = [clean(cell.text_content()) for cell in tr.xpath("./th|./td")]
                if cells:
                    rows.append(cells)
            if rows:
                lens = [max(len(row[i]) for row in rows if i < len(row)) for i in range(len(rows[0]))]
                add_table(doc, rows, [max(1, min(5, n / 14)) for n in lens], caption)
        elif tag == "div":
            classes = set((node.get("class") or "").split())
            if "formula" in classes:
                code = node.find(".//code")
                small = node.find(".//small")
                add_formula(doc, code.text_content() if code is not None else node.text_content(), small.text_content() if small is not None else "")
            elif "process-strip" in classes:
                steps = []
                for item in node.xpath("./div"):
                    title = item.find("b")
                    title_text = clean(title.text_content()) if title is not None else "Bước"
                    full = clean(item.text_content())
                    steps.append([title_text, clean(full[len(title_text):])])
                if steps:
                    add_table(doc, [[s[0] for s in steps], [s[1] for s in steps]], [1] * len(steps), "Quy trình tổng quát")
            elif "callout" in classes:
                strong = node.find(".//strong")
                label = strong.text_content() if strong is not None else "Lưu ý"
                text = clean(node.text_content())
                add_callout(doc, label, text[len(clean(label)):].strip())
            else:
                text = clean(node.text_content())
                if text:
                    doc.add_paragraph(text)


def add_manual_toc(doc, instructor=False):
    doc.add_heading("MỤC LỤC", level=1)
    if not instructor:
        for part in COURSE["parts"]:
            p = doc.add_paragraph()
            set_run(p.add_run(clean(part["title"])), bold=True, color=DARK_BLUE)
            for ch in [x for x in COURSE["chapters"] if x["part"] == part["id"]]:
                item = doc.add_paragraph()
                item.paragraph_format.left_indent = Inches(0.18)
                item.paragraph_format.space_after = Pt(3)
                set_run(item.add_run(f"Chương {ch['number']}. {clean(ch['title'])}"), color=INK)
    else:
        add_numbered_items(doc, ["Thiết kế học phần và ma trận CLO", "Kế hoạch 15 tuần", "Hướng dẫn giảng dạy 15 chương", "Ngân hàng đáp án", "Rubric và biểu mẫu", "Chính sách AI và kiểm soát chất lượng"])
    doc.add_page_break()


def add_front_matter(doc):
    doc.add_heading("LỜI NÓI ĐẦU", level=1)
    doc.add_paragraph("Giáo trình này được xây dựng cho đào tạo liên ngành giữa khoa học môi trường, dữ liệu, GIS và trí tuệ nhân tạo. Mục tiêu không phải biến mọi vấn đề thành bài toán AI, mà giúp người học chọn đúng ranh giới, bảo toàn đơn vị, xây bằng chứng có nguồn gốc và sử dụng mô hình đúng vai trò trong quyết định.")
    doc.add_paragraph("Ấn bản 2.0 mở rộng bản website ban đầu bằng các chuyên đề Digital Earth, PINNs, GNN, foundation models, trung hòa carbon, công bằng môi trường và Digital Twin lưu vực. Các góp ý do người dùng cung cấp được tiếp nhận theo chủ đề và đối chiếu bằng nguồn học thuật/chính thức; tên cá nhân hoặc tổ chức chưa được xác minh không được xem là chứng thực.")
    add_callout(doc, "Định vị Mô hình LONG", "LONG là khung sư phạm do tác giả đề xuất. Khi dùng LCA, EIA, DPSIR, MCA, mô hình thủy văn hoặc ML, người học vẫn phải tuân thủ tên gọi, giả định và quy trình kiểm định của phương pháp tương ứng.")
    doc.add_heading("CÁCH SỬ DỤNG GIÁO TRÌNH", level=1)
    add_numbered_items(doc, [
        "Đọc tổng quan và chuẩn đầu ra trước mỗi chương.",
        "Giải thích mọi công thức bằng ranh giới, đơn vị và điều kiện áp dụng.",
        "Thực hiện bài thực hành bằng dữ liệu thật có giấy phép hoặc dữ liệu giả định được gắn nhãn.",
        "Làm tự kiểm tra trước khi xem đáp án ở phụ lục.",
        "Lưu dữ liệu, mã, tham số và nhật ký để người khác có thể tái lập.",
    ])
    doc.add_picture(str(ASSETS / "long_framework.png"), width=Inches(6.45))
    set_last_image_alt(doc, "Sơ đồ bốn trụ LONG: Learning và vòng đời, Optimization, Network và Governance")
    cap = doc.add_paragraph("Hình 0.1. Bốn trụ của khung sư phạm LONG.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cap.runs[0], size=9, italic=True, color=MUTED)


def add_chapter(doc, chapter, ref_map):
    start_on_new_page(doc.add_heading(f"CHƯƠNG {chapter['number']}. {clean(chapter['title']).upper()}", level=1))
    p = doc.add_paragraph()
    set_run(p.add_run(clean(chapter["summary"])), italic=True, color=MUTED)
    add_callout(doc, "Thời lượng và mức độ", f"{chapter['duration']} | {chapter['level']} | Từ khóa: {', '.join(chapter['keywords'])}")
    doc.add_heading("Chuẩn đầu ra chương", level=2)
    add_numbered_items(doc, chapter["outcomes"])

    if chapter["number"] == 1:
        doc.add_picture(str(ASSETS / "ipof_loop.png"), width=Inches(6.45))
        set_last_image_alt(doc, "Vòng Đầu vào, Tiến trình, Đầu ra và Phản hồi của hệ thống môi trường")
        p = doc.add_paragraph("Hình 1.1. Vòng Đầu vào - Tiến trình - Đầu ra - Phản hồi.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.runs[0], size=9, italic=True, color=MUTED)
    if chapter["number"] == 14:
        doc.add_picture(str(ASSETS / "digital_twin_architecture.png"), width=Inches(6.45))
        set_last_image_alt(doc, "Kiến trúc Digital Twin lưu vực gồm quan trắc, nền dữ liệu, lõi mô hình và quyết định")
        p = doc.add_paragraph("Hình 14.1. Kiến trúc tham chiếu Digital Twin lưu vực.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.runs[0], size=9, italic=True, color=MUTED)

    for section in chapter["sections"]:
        doc.add_heading(clean(section["title"]), level=2)
        html_fragment_to_doc(doc, section["html"])

    doc.add_heading("Ca nghiên cứu", level=2)
    add_callout(doc, chapter["caseStudy"]["title"], chapter["caseStudy"]["body"])
    doc.add_heading("Thực hành có hướng dẫn", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run(f"{chapter['lab']['code']} - {clean(chapter['lab']['title'])}"), bold=True, color=DARK_BLUE)
    add_numbered_items(doc, chapter["lab"]["tasks"])

    doc.add_heading("Bài tập mở rộng", level=2)
    add_numbered_items(doc, [ex["title"] for ex in chapter["exercises"]])
    doc.add_heading("Tự kiểm tra", level=2)
    for idx, quiz in enumerate(chapter["quiz"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        set_run(p.add_run(f"Câu {idx}. {clean(quiz['q'])}"), bold=True, color=INK)
        add_bullet_items(doc, [f"{letter}. {clean(option)}" for letter, option in zip("ABCD", quiz["options"])], size=9.5)

    doc.add_heading("Nguồn học tập chính", level=2)
    ref_paragraphs = []
    for rid in chapter.get("refs", []):
        ref = ref_map.get(rid)
        if not ref:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        ref_paragraphs.append(p)
        set_run(p.add_run(clean(ref["citation"]) + " "), size=8.5, color=INK)
        add_hyperlink(p, "Truy cập nguồn", ref["url"])
    if ref_paragraphs:
        num_id = create_numbering(doc, bullet=True)
        for p in ref_paragraphs:
            num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = 0
            num_pr.get_or_add_numId().val = num_id


def add_answer_key(doc, *, page_break_before=False):
    heading = doc.add_heading("PHỤ LỤC A. ĐÁP ÁN VÀ GỢI Ý", level=1)
    if page_break_before:
        start_on_new_page(heading)
    for chapter in COURSE["chapters"]:
        doc.add_heading(f"Chương {chapter['number']}", level=2)
        for idx, quiz in enumerate(chapter["quiz"], 1):
            answer = "ABCD"[quiz["answer"]]
            p = doc.add_paragraph()
            set_run(p.add_run(f"Câu {idx}: {answer}. "), bold=True, color=DARK_BLUE)
            set_run(p.add_run(clean(quiz["explain"])), color=INK)
        for idx, ex in enumerate(chapter["exercises"], 1):
            p = doc.add_paragraph()
            set_run(p.add_run(f"Bài tập {idx} - Gợi ý: "), bold=True, color=DARK_BLUE)
            set_run(p.add_run(clean(ex["solution"])), color=INK)


def build_student_book():
    doc = Document()
    configure_document(doc, "narrative_proposal", "PHÂN TÍCH HỆ THỐNG MÔI TRƯỜNG | ẤN BẢN 2.0")
    add_cover(doc)
    add_manual_toc(doc)
    add_front_matter(doc)
    ref_map = {r["id"]: r for r in COURSE["references"]}
    current_part = None
    for chapter in COURSE["chapters"]:
        if current_part != chapter["part"]:
            part = next(p for p in COURSE["parts"] if p["id"] == chapter["part"])
            start_on_new_page(doc.add_heading(clean(part["title"]).upper(), level=1))
            doc.add_paragraph("Phần này tổ chức kiến thức theo tiến trình năng lực và kết nối trực tiếp với thực hành, tự kiểm tra và đồ án.")
            current_part = chapter["part"]
        add_chapter(doc, chapter, ref_map)
    add_answer_key(doc, page_break_before=True)
    start_on_new_page(doc.add_heading("PHỤ LỤC B. RUBRIC ĐỒ ÁN", level=1))
    add_table(doc, [
        ["Tiêu chí", "Tỷ trọng", "Minh chứng đạt"],
        ["Câu hỏi, ranh giới, bên liên quan", "15%", "Quyết định rõ, biên có lý do, nhu cầu người dùng"],
        ["Dữ liệu, QA/QC, provenance", "20%", "Schema, nguồn, cờ chất lượng, quyền sử dụng"],
        ["Mô hình, kiểm định, bất định", "25%", "Baseline, test ngoài mẫu, phần dư, khoảng dự báo"],
        ["Phân tích LONG", "20%", "Ít nhất hai trụ tạo giá trị quyết định"],
        ["Sản phẩm, đạo đức, tái lập", "20%", "Mã, model/data card, policy brief, fallback"],
    ], [2.2, .7, 3.6], "Bảng B.1. Rubric đồ án tích hợp")
    start_on_new_page(doc.add_heading("PHỤ LỤC C. TÀI LIỆU THAM KHẢO", level=1))
    ref_paragraphs = []
    for ref in COURSE["references"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        ref_paragraphs.append(p)
        set_run(p.add_run(clean(ref["citation"]) + " "), size=8.5)
        add_hyperlink(p, "Truy cập nguồn", ref["url"])
    num_id = create_numbering(doc, bullet=True)
    for p in ref_paragraphs:
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
    start_on_new_page(doc.add_heading("PHỤ LỤC D. BÁO CÁO KIỂM SOÁT CHẤT LƯỢNG", level=1))
    add_table(doc, [
        ["Vòng", "Phạm vi", "Kết quả"],
        ["1", "Cấu trúc và sư phạm", "15 chương; chuẩn đầu ra - lý thuyết - case - thực hành - quiz"],
        ["2", "Khoa học và đơn vị", "Công thức kèm giả định, đơn vị và giới hạn"],
        ["3", "Nguồn, đạo đức, chứng thực", "Ưu tiên nguồn chính thức; không coi thư chưa xác minh là chứng thực"],
        ["4", "Khả dụng số", "Website, Word, PDF, EPUB, PWA, responsive và liên kết chương"],
    ], [0.6, 2.3, 3.6], "Bảng D.1. Tổng hợp vòng kiểm soát")
    out = DOWNLOADS / "Giao_trinh_AI_LONG.docx"
    doc.save(out)
    return out


def build_instructor_manual():
    doc = Document()
    configure_document(doc, "compact_reference_guide", "SỔ TAY GIẢNG VIÊN | AI × HỆ THỐNG MÔI TRƯỜNG")
    add_cover(doc, instructor=True)
    add_manual_toc(doc, instructor=True)
    doc.add_heading("1. THIẾT KẾ HỌC PHẦN", level=1)
    doc.add_paragraph("Học phần 4 tín chỉ sử dụng flipped learning, CDIO và project-based learning. Giảng viên đánh giá khả năng đặt ranh giới, kiểm tra đơn vị, thiết kế kiểm định và truyền đạt bất định; không chấm cao cho mô hình phức tạp nếu thiếu baseline hoặc trách nhiệm vận hành.")
    add_table(doc, [
        ["Thành phần", "Tỷ trọng", "Sản phẩm"],
        ["Thực hành cá nhân", "30%", "Notebook, sơ đồ, bảng tính, model card"],
        ["Dự án nhóm", "30%", "Mô hình, dashboard/twin, dữ liệu, policy brief"],
        ["Thuyết trình/phản biện", "20%", "Seminar, review đồng cấp, bảo vệ lựa chọn"],
        ["Cuối kỳ mở", "20%", "Phân tích tình huống và lập luận"],
    ], [2.1, .8, 3.6], "Bảng 1.1. Cơ cấu đánh giá")
    doc.add_heading("2. KẾ HOẠCH 15 TUẦN", level=1)
    week_rows = [["Tuần", "Chương", "Hoạt động trước lớp", "Sản phẩm trên lớp"]]
    plan = [
        (1, "1", "Đọc IPOF; quiz", "Bản đồ hệ thống"), (2, "2", "Ôn cân bằng", "Mô hình BOD"),
        (3, "3", "Kiểm tra metadata", "Data dictionary"), (4, "4", "Baseline ML", "Sơ đồ chia tập"),
        (5, "4", "Đọc PINN/GNN/FM", "Model critique"), (6, "5", "Ranh giới LCA", "Mini-LCA"),
        (7, "6", "Pareto", "Bài toán tối ưu"), (8, "7", "Đồ thị", "Mạng đa lớp"),
        (9, "8", "Rủi ro/DSS", "Protocol cảnh báo"), (10, "9-10", "Pilot/biodiversity", "Thiết kế quan trắc"),
        (11, "11", "Risk map", "Dashboard đô thị"), (12, "12", "GHG/MRV", "MACC"),
        (13, "13", "Environmental justice", "Bản đồ công bằng"), (14, "14", "Twin maturity", "Kiến trúc twin"),
        (15, "15", "Hoàn thiện hồ sơ", "Bảo vệ capstone"),
    ]
    for row in plan:
        week_rows.append([str(x) for x in row])
    add_table(doc, week_rows, [.6, .8, 2.4, 2.7], "Bảng 2.1. Kế hoạch giảng dạy")

    doc.add_heading("3. HƯỚNG DẪN GIẢNG DẠY THEO CHƯƠNG", level=1)
    misconceptions = {
        1: "Cộng trực tiếp nồng độ; hiểu phản hồi dương là tốt.", 2: "Hiệu chỉnh và kiểm định trên cùng dữ liệu.",
        3: "Xóa mọi ngoại lai; coi FAIR là công khai hoàn toàn.", 4: "SHAP là nhân quả; mô hình sâu luôn tốt hơn.",
        5: "Thay đổi ranh giới LCA mà không báo; đếm trùng lợi ích.", 6: "Biến mọi tiêu chí thành một điểm số tùy ý.",
        7: "Cạnh tương quan được diễn giải như dòng vật chất.", 8: "Dashboard được coi là DSS dù không có hành động.",
        9: "Tối ưu hiệu suất nồng độ dù lưu lượng thay đổi.", 10: "Tỷ lệ phát hiện bằng độ phong phú.",
        11: "Hazard map được gọi là risk map.", 12: "AI thay MRV; baseline không ghi giả định.",
        13: "Truy xuất tự bảo đảm công bằng.", 14: "Mô hình 3D tĩnh được gọi là twin.",
        15: "Chỉ bàn giao PDF/ảnh dashboard, thiếu dữ liệu và mã.",
    }
    for ch in COURSE["chapters"]:
        doc.add_heading(f"3.{ch['number']}. Chương {ch['number']} - {clean(ch['shortTitle'])}", level=2)
        add_callout(doc, "Trọng tâm", ch["summary"])
        p = doc.add_paragraph()
        set_run(p.add_run("Ngộ nhận cần xử lý: "), bold=True, color=DARK_BLUE)
        set_run(p.add_run(misconceptions[ch["number"]]), color=INK)
        p = doc.add_paragraph()
        set_run(p.add_run("Sản phẩm đánh giá: "), bold=True, color=DARK_BLUE)
        set_run(p.add_run(f"{ch['lab']['code']} - {clean(ch['lab']['title'])}"), color=INK)
        add_bullet_items(doc, ch["lab"]["tasks"])
        p = doc.add_paragraph()
        set_run(p.add_run("Câu hỏi phản biện mở: "), bold=True, color=DARK_BLUE)
        set_run(p.add_run(f"Điều kiện nào làm kết luận của chương {ch['number']} không còn đáng tin cậy hoặc không còn công bằng?"), color=INK)

    doc.add_heading("4. NGÂN HÀNG ĐÁP ÁN", level=1)
    add_answer_key(doc)
    doc.add_heading("5. RUBRIC VÀ BIỂU MẪU", level=1)
    add_table(doc, [
        ["Mức", "Mô tả tổng quát"],
        ["4 - Xuất sắc", "Đúng khoa học; tái lập; bất định rõ; LONG tạo giá trị; có audit/fallback"],
        ["3 - Đạt tốt", "Phương pháp đúng; còn thiếu một số kiểm tra độ nhạy hoặc tài liệu hóa"],
        ["2 - Đạt tối thiểu", "Có sản phẩm nhưng ranh giới/kiểm định/đạo đức chưa đầy đủ"],
        ["1 - Chưa đạt", "Sai đơn vị hoặc rò rỉ; thiếu nguồn; kết luận quá mức; không thể tái lập"],
    ], [1.2, 5.3], "Bảng 5.1. Thang mô tả chung")
    doc.add_heading("6. CHÍNH SÁCH SỬ DỤNG AI", level=1)
    add_bullet_items(doc, [
        "Sinh viên được dùng AI để gợi ý mã, cấu trúc và phản biện, nhưng phải công bố phạm vi sử dụng.",
        "Không chấp nhận trích dẫn, dữ liệu hoặc kết quả mô hình do AI tạo mà chưa kiểm chứng.",
        "Người học chịu trách nhiệm về đơn vị, dữ liệu, phép tính, bản quyền và kết luận.",
        "Bài nộp phải gồm nhật ký quyết định; giảng viên có thể yêu cầu giải thích trực tiếp bất kỳ đoạn mã hoặc kết quả nào.",
        "Dữ liệu nhạy cảm không được tải lên dịch vụ AI khi chưa có thẩm quyền và biện pháp bảo vệ.",
    ], size=9.5)
    out = DOWNLOADS / "So_tay_giang_vien_AI_LONG.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    make_diagrams()
    student = build_student_book()
    instructor = build_instructor_manual()
    print(student)
    print(instructor)
